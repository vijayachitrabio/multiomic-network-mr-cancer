#!/usr/bin/env python3
"""
Script 44: Generate professional project-summary email as Word document
        with 4 embedded figures.
Output: PROJECT_SUMMARY_EMAIL_2026-05-08.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, copy

proj = "/Users/vijayachitramodhukur/Library/Mobile Documents/com~apple~CloudDocs/ECLAI/MultiOmic_Network_MR_Project"
fig_dir = os.path.join(proj, "results", "figures")
out_path = os.path.join(proj, "PROJECT_SUMMARY_EMAIL_2026-05-08.docx")

FIGURES = [
    {
        "file": "fig1_phase2_forest.png",
        "label": "Figure 1",
        "title": "Proteome-wide MR screen — forest plot of 17 protein–cancer associations",
        "caption": (
            "Forest plot showing odds ratios (95% CI) for all 17 FDR-significant protein–cancer "
            "associations from the proteome-wide MR screen (701 proteins, FinnGen Olink pQTL, N = 619). "
            "Proteins are colour-coded by evidence tier. Tier 1 proteins (EFNA1, TNFRSF6B, ATRAID) are "
            "supported by both MR and strong SuSiE colocalization (PPH4 ≥ 0.88). Tier 2a proteins "
            "(SNX15, PM20D1) are corroborated by Bonferroni-significant MAGMA gene-level evidence. "
            "The dashed vertical line marks the null (OR = 1.0)."
        ),
        "width": Inches(5.8),
    },
    {
        "file": "fig3_triangulation.png",
        "label": "Figure 2",
        "title": "Multi-layer evidence triangulation across five independent frameworks",
        "caption": (
            "Heatmap summarising evidence across five independent analytical layers for the 17 "
            "MR-prioritised proteins: MR effect estimate and FDR, SuSiE colocalization posterior "
            "(PPH4), MAGMA gene-level p-value, ER-subtype specificity, and mediation indirect effect. "
            "Tier assignments are shown on the left margin. The dramatic coloc.abf vs coloc.susie "
            "discordance for EFNA1 and ATRAID — both loci harbouring nine independent GWAS credible "
            "sets — is annotated."
        ),
        "width": Inches(5.8),
    },
    {
        "file": "fig5_metabolite_cancer_coloc.png",
        "label": "Figure 3",
        "title": "GCKR-centred metabolic susceptibility architecture — metabolite–cancer colocalization",
        "caption": (
            "Locus-by-locus colocalization results across 13,100 genomic windows for four FDR-significant "
            "metabolite–breast cancer pairs (glycine, total BCAA, TG:PG ratio, HDL-C). The GCKR locus "
            "(rs1260326 P446L, chr2:27 Mb, GRCh38) shows near-certain colocalization with breast cancer "
            "for glycine, total BCAA, and TG:PG (PPH4 ≈ 1.000). Secondary colocalization loci at CPS1 "
            "(chr2:211 Mb), MLXIPL/ChREBP (chr7:72 Mb), and MC4R (chr16) extend the metabolic "
            "susceptibility map."
        ),
        "width": Inches(5.8),
    },
    {
        "file": "fig4_mediation_paths.png",
        "label": "Figure 4",
        "title": "Two-step mediation MR — protein–metabolite–breast cancer indirect effect paths",
        "caption": (
            "Results of two-step mediation MR for six protein–metabolite–breast cancer triplets. "
            "Indirect effects were estimated by the product-of-coefficients method with delta-method "
            "standard errors. Step-2 MR (metabolite → cancer) used both IVW and weighted median (WM) "
            "estimators; WM is shown as the primary result given GCKR pleiotropy inflating IVW for "
            "Total_BCAA paths. Five of six paths were significant under WM; the ATRAID → TG:PG path "
            "was rejected (WM p = 0.627). Proportion mediated ranged from 2.2% (ITIH3–glycine) to "
            "16.1% (IL34–Total BCAA)."
        ),
        "width": Inches(5.8),
    },
]

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(doc, text, style="Normal", align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.alignment = align
    run = p.add_run(text)
    set_font(run)
    return p

def add_heading(doc, text, level=1, space_before=12, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    sizes = {1: 14, 2: 12, 3: 11}
    set_font(run, size=sizes.get(level, 11), bold=True,
             color=(31, 73, 125) if level == 1 else (0, 70, 127) if level == 2 else (0, 0, 0))
    return p

def add_bullet(doc, text, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
    run = p.add_run(text)
    set_font(run)
    return p

def add_figure(doc, fig_dict):
    """Insert figure image + bold label + caption paragraph."""
    # Figure label
    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(10)
    lp.paragraph_format.space_after  = Pt(2)
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = lp.add_run(f"{fig_dict['label']}. ")
    set_font(lr, bold=True, size=10)
    tr = lp.add_run(fig_dict['title'])
    set_font(tr, bold=True, italic=True, size=10)

    # Image
    img_path = os.path.join(fig_dir, fig_dict["file"])
    ip = doc.add_paragraph()
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.paragraph_format.space_before = Pt(2)
    ip.paragraph_format.space_after  = Pt(2)
    run = ip.add_run()
    run.add_picture(img_path, width=fig_dict["width"])

    # Caption
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(4)
    cp.paragraph_format.space_after  = Pt(16)
    cp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cr = cp.add_run(fig_dict["caption"])
    set_font(cr, size=9, color=(89, 89, 89))

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4472C4')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ── Build document ────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── HEADER BLOCK ─────────────────────────────────────────────────────────────
hp = doc.add_paragraph()
hp.paragraph_format.space_after = Pt(2)
hr = hp.add_run("MultiOmic Network MR Project — Analysis Summary & Future Work")
set_font(hr, size=16, bold=True, color=(31, 73, 125))

dp = doc.add_paragraph()
dp.paragraph_format.space_after = Pt(2)
dr = dp.add_run("Date: 8 May 2026    |    Version: v0.4    |    Target journal: BMC Medicine")
set_font(dr, size=10, color=(128, 128, 128))

add_divider(doc)
doc.add_paragraph()

# ── EMAIL HEADER ─────────────────────────────────────────────────────────────
for line, bold in [
    ("To:      [Collaborator name / PI]", False),
    ("From:    Vijaya Chitramodi Modhukur", False),
    ("Subject: MultiOmic Network MR — Analysis Summary & Next Steps (v0.4 Manuscript Ready)", True),
    ("Date:    8 May 2026", False),
]:
    ep = doc.add_paragraph()
    ep.paragraph_format.space_after = Pt(1)
    er = ep.add_run(line)
    set_font(er, size=10, bold=bold, color=(0,0,0) if bold else (64,64,64))

doc.add_paragraph()
add_divider(doc)

# ── SALUTATION ───────────────────────────────────────────────────────────────
add_para(doc, "Dear [Collaborator],", space_before=10, space_after=6)
add_para(doc,
    "I wanted to share a concise summary of where we stand with the proteome-wide Mendelian "
    "randomisation project ahead of next week's finalisation push. Please find the four key "
    "figures embedded below, with a brief description of each analytical contribution.",
    space_before=0, space_after=10)

# ── SECTION 1: WHAT HAS BEEN COMPLETED ───────────────────────────────────────
add_heading(doc, "1. What has been completed", level=1)

add_heading(doc, "MR screen", level=2, space_before=6, space_after=2)
add_para(doc,
    "A proteome-wide two-sample MR screen of 701 circulating proteins (FinnGen Olink pQTL, N = 619) "
    "was conducted against breast, endometrial, and ovarian cancer GWAS. Seventeen protein–cancer "
    "associations survived 5% FDR correction: 16 for breast cancer and one (ABO) for endometrial cancer; "
    "no ovarian signal survived correction. All instruments passed Steiger directionality filtering — "
    "zero reverse-causation signals detected.", space_after=6)

add_heading(doc, "Protein colocalization (coloc.abf + SuSiE)", level=2, space_before=6, space_after=2)
add_para(doc,
    "Colocalization was completed for all 8 priority proteins using both coloc.abf and SuSiE-based "
    "coloc.susie with 1000 Genomes EUR LD (N = 633). A key methodological finding: two proteins "
    "classified as distinct-causal-variant loci by coloc.abf (PPH3 > 0.90) were correctly reclassified "
    "as strongly colocalized by coloc.susie — because their loci each harbour nine independent GWAS "
    "credible sets that coloc.abf cannot handle.", space_after=4)
add_bullet(doc, "EFNA1: coloc.abf PPH3 = 0.901  →  coloc.susie PPH4 = 0.963  [Tier 1]")
add_bullet(doc, "ATRAID: coloc.abf PPH3 = 0.997  →  coloc.susie PPH4 = 0.996  [Tier 1]")
add_bullet(doc, "TNFRSF6B: concordant strong colocalization by both methods (PPH4 ≈ 0.885)  [Tier 1]")
doc.add_paragraph()

add_heading(doc, "MAGMA gene-level triangulation", level=2, space_before=6, space_after=2)
add_para(doc,
    "Gene-based association testing (MAGMA, breast GWAS N = 228,951, 17,545 genes) provided "
    "an independent instrument-free line of evidence. SNX15 (p = 1.47×10⁻¹¹, rank 109/17,545) "
    "and PM20D1 (p = 1.43×10⁻⁶) reached Bonferroni significance — classified as Tier 2a. "
    "10 of 16 breast cancer proteins showed nominal MAGMA enrichment (binomial p = 5.9×10⁻¹⁰ "
    "vs null expectation).", space_after=6)

add_heading(doc, "Metabolite–cancer colocalization", level=2, space_before=6, space_after=2)
add_para(doc,
    "Exhaustive locus-by-locus colocalization across 13,100 genomic windows confirmed GCKR "
    "rs1260326 (P446L, chr2:27 Mb) as the dominant shared locus for amino-acid and lipid "
    "metabolism effects on breast cancer (PPH4 ≈ 1.000 for glycine, total BCAA, and TG:PG ratio). "
    "Secondary loci at CPS1, MLXIPL/ChREBP, MC4R, and near VEGFA extend the metabolic "
    "susceptibility architecture.", space_after=6)

add_heading(doc, "Two-step mediation MR", level=2, space_before=6, space_after=2)
add_para(doc,
    "Six protein–metabolite–breast cancer triplets were tested. Five paths showed significant "
    "indirect effects under a pleiotropy-robust weighted median (WM) estimator. The ATRAID → TG:PG "
    "path was correctly rejected (WM p = 0.627) after the IVW result was found to reflect GCKR "
    "pleiotropy.", space_after=4)
add_bullet(doc, "IL34 → Total BCAA → Breast: WM p_indirect = 0.0014, proportion mediated = 16.1%")
add_bullet(doc, "TNFRSF6B → Total BCAA → Breast: WM p_indirect = 0.022, proportion mediated = 9.0%")
add_bullet(doc, "APOE → Glycine → Breast: WM p_indirect = 0.00035, proportion mediated = 4.3%")
add_bullet(doc, "EFNA1 → Total BCAA → Breast: WM p_indirect = 0.018, proportion mediated = 3.1%")
add_bullet(doc, "ITIH3 → Glycine → Breast: WM p_indirect = 0.031, proportion mediated = 2.2%")
doc.add_paragraph()

add_heading(doc, "Druggability and pathways", level=2, space_before=6, space_after=2)
add_para(doc,
    "OpenTargets assessment flagged FGFR4 (17 drugs; orantinib Phase 3) and KLB (fazpilodemab "
    "Phase 2) as having existing clinical programmes. The remaining 15 proteins — including all "
    "three Tier 1 candidates — have no registered drug programme, representing novel target "
    "hypotheses. g:Profiler enrichment (background = 701 proteins) identified 38 significant "
    "terms dominated by FGFR4–KLB signalling and IGF1R/PI3K/AKT/MAPK pathways.", space_after=6)

add_para(doc,
    "The updated manuscript (Results + Discussion, v0.4, dated 2026-05-08) is written to BMC "
    "Medicine standard and is ready for your review. Remaining [PLACEHOLDER] and [REF] tags mark "
    "sections pending next week's analyses.", space_after=10)

# ── SECTION 2: REMAINING WORK ─────────────────────────────────────────────────
add_heading(doc, "2. Remaining work before submission", level=1)

add_heading(doc, "High priority", level=2, space_before=6, space_after=2)
add_bullet(doc, "Bidirectional MR (cancer liability → protein levels) — standard reviewer requirement")
add_bullet(doc, "MVMR conditioning on BMI and CRP for the three Tier 1 proteins (EFNA1, TNFRSF6B, ATRAID)")
add_bullet(doc, "deCODE replication (N = 35,559) for SNX15 and EFNA1 — scripts ready, access token needed")

add_heading(doc, "Medium priority", level=2, space_before=6, space_after=2)
add_bullet(doc, "Fill numerical placeholders (instrument F-statistic range, top pathway enrichment terms)")
add_bullet(doc, "Add PMIDs / DOIs to all [REF] tags throughout Results and Discussion")
add_bullet(doc, "Formal colocalization for the nine Tier 2d proteins (currently MR-only)")

add_heading(doc, "Lower priority", level=2, space_before=6, space_after=2)
add_bullet(doc, "Full step-1 mediation sensitivity for remaining proteins (FGF23, RANKL, PM20D1, KLB, ABO, INHBB)")
add_bullet(doc, "Resolve FGFR4 directional paradox (protective MR vs inhibitor development) experimentally")
doc.add_paragraph()

# ── SECTION 3: FIGURES ────────────────────────────────────────────────────────
add_heading(doc, "3. Key figures", level=1)
add_para(doc,
    "Four figures are embedded below, covering the four analytical stories of the paper: "
    "MR discovery, colocalization methodology, metabolic architecture, and mediation mechanism.",
    space_after=10)

for fig in FIGURES:
    add_figure(doc, fig)

# ── SIGN-OFF ──────────────────────────────────────────────────────────────────
add_divider(doc)
doc.add_paragraph()
add_para(doc, "Happy to discuss any of the above — looking forward to finalising the submission next week.",
         space_after=4)
add_para(doc, "Best wishes,", space_after=2)
add_para(doc, "Vijaya Chitramodi Modhukur", space_after=1)
p = doc.add_paragraph()
r = p.add_run("vijayachitramodhukur@[institution.edu]  |  MultiOmic Network MR Project  |  2026-05-08")
set_font(r, size=9, color=(128,128,128))

# ── SAVE ─────────────────────────────────────────────────────────────────────
doc.save(out_path)
print(f"✓ Saved → {out_path}")
